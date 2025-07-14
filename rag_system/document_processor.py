import os
import PyPDF2
from docx import Document
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import streamlit as st


class DocumentProcessor:
    """Handles document loading and processing for the RAG system."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
        )

    def load_pdf(self, file_path: str) -> str:
        """Load text from a PDF file."""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            st.error(f"Error loading PDF {file_path}: {str(e)}")
            return ""

    def load_docx(self, file_path: str) -> str:
        """Load text from a DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error loading DOCX {file_path}: {str(e)}")
            return ""

    def load_text(self, file_path: str) -> str:
        """Load text from a text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except Exception as e:
            st.error(f"Error loading text file {file_path}: {str(e)}")
            return ""

    def process_file(
        self, file_path: str, original_name: str = None
    ) -> List[LangchainDocument]:
        """Process a single file and return chunks as LangchainDocument objects."""
        file_extension = os.path.splitext(file_path)[1].lower()

        # Load text based on file type
        if file_extension == ".pdf":
            text = self.load_pdf(file_path)
        elif file_extension == ".docx":
            text = self.load_docx(file_path)
        elif file_extension in [".txt", ".md"]:
            text = self.load_text(file_path)
        else:
            st.warning(f"Unsupported file type: {file_extension}")
            return []

        if not text.strip():
            st.warning(f"No text extracted from {file_path}")
            return []

        # Split text into chunks
        chunks = self.text_splitter.split_text(text)

        # Use original filename if provided, otherwise use the file path
        source_name = original_name if original_name else file_path

        # Convert to LangchainDocument objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": source_name,
                    "chunk_id": i,
                    "file_type": file_extension,
                },
            )
            documents.append(doc)

        return documents

    def process_files(
        self, file_paths: List[str], original_names: List[str] = None
    ) -> List[LangchainDocument]:
        """Process multiple files and return all chunks."""
        all_documents = []

        for i, file_path in enumerate(file_paths):
            original_name = (
                original_names[i]
                if original_names and i < len(original_names)
                else None
            )
            documents = self.process_file(file_path, original_name)
            all_documents.extend(documents)

        return all_documents

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get information about a file."""
        file_size = os.path.getsize(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()

        return {
            "name": os.path.basename(file_path),
            "size": file_size,
            "extension": file_extension,
            "path": file_path,
        }
